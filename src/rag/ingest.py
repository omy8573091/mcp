from __future__ import annotations

import io
import os
from typing import Iterable, List, Tuple
import hashlib

import pandas as pd
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from opentelemetry import trace

from .db import db_session
from .models import Base, Document, Chunk, EmbeddingCache
from .openai_utils import embed_texts
from .settings import get_rag_settings

tracer = trace.get_tracer(__name__)


def _chunk_text(text: str, max_chars: int = 1200, overlap: int = 100) -> List[str]:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    chunks: List[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + max_chars)
        chunks.append(text[start:end])
        if end == len(text):
            break
        start = end - overlap
        if start < 0:
            start = 0
    return chunks


def _read_pdf(data: bytes) -> str:
    with fitz.open(stream=data, filetype="pdf") as doc:
        texts = [page.get_text() for page in doc]
    return "\n".join(texts)


def _read_docx(data: bytes) -> str:
    with io.BytesIO(data) as bio:
        d = DocxDocument(bio)
        return "\n".join(p.text for p in d.paragraphs)


def _read_text(data: bytes) -> str:
    return data.decode(errors="ignore")


def _read_csv_or_xlsx(data: bytes, content_type: str) -> str:
    with io.BytesIO(data) as bio:
        if content_type in {"text/csv", "csv"}:
            df = pd.read_csv(bio)
        else:
            df = pd.read_excel(bio)
    return df.to_string(index=False)


def _detect_type(filename: str) -> str:
    ext = os.path.splitext(filename.lower())[1]
    if ext == ".pdf":
        return "application/pdf"
    if ext in {".docx"}:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if ext in {".txt", ".log", ".md"}:
        return "text/plain"
    if ext in {".csv"}:
        return "text/csv"
    if ext in {".xls", ".xlsx"}:
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    return "application/octet-stream"


def ingest_file(filename: str, data: bytes, *, source_path: str | None = None) -> int:
    with tracer.start_as_current_span("rag.ingest_file") as span:
        span.set_attributes({
            "rag.filename": filename,
            "rag.file_size_bytes": len(data),
            "rag.source_path": source_path or "",
        })
        
        content_type = _detect_type(filename)
        span.set_attribute("rag.content_type", content_type)
        
        with tracer.start_as_current_span("rag.parse_content"):
            if content_type == "application/pdf":
                text = _read_pdf(data)
            elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = _read_docx(data)
            elif content_type in {"text/plain"}:
                text = _read_text(data)
            elif content_type in {"text/csv", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"}:
                text = _read_csv_or_xlsx(data, "text/csv" if content_type == "text/csv" else "xlsx")
            else:
                text = _read_text(data)
            
            span.set_attribute("rag.text_length", len(text))

        with tracer.start_as_current_span("rag.chunk_text"):
            chunks = _chunk_text(text)
            span.set_attribute("rag.chunk_count", len(chunks))

        settings = get_rag_settings()
        embeddings: List[List[float]] = []
        
        with tracer.start_as_current_span("rag.embed_chunks"):
            if settings.embed_cache_enable:
                with db_session() as s:
                    from sqlalchemy import select

                    for ch in chunks:
                        sha = hashlib.sha256(ch.encode()).hexdigest()
                        row = s.execute(select(EmbeddingCache).where(EmbeddingCache.sha256 == sha)).scalar_one_or_none()
                        if row:
                            embeddings.append(row.embedding)
                        else:
                            embeddings.append([])  # placeholder, fill after embedding
                    # compute for missing
                    to_embed = [chunks[i] for i, e in enumerate(embeddings) if not e]
                    if to_embed:
                        new_embs = embed_texts(to_embed)
                        it = iter(new_embs)
                        for i, e in enumerate(embeddings):
                            if not e:
                                emb = next(it)
                                embeddings[i] = emb
                                sha = hashlib.sha256(chunks[i].encode()).hexdigest()
                                s.add(EmbeddingCache(sha256=sha, embedding=emb))
                    s.flush()
                    span.set_attribute("rag.cache_hits", len(embeddings) - len(to_embed))
                    span.set_attribute("rag.cache_misses", len(to_embed))
            else:
                embeddings = embed_texts(chunks)
                span.set_attribute("rag.cache_enabled", False)

        with tracer.start_as_current_span("rag.save_document"):
            with db_session() as s:
                from sqlalchemy import text as sql_text

                # Ensure extensions/tables exist
                s.execute(sql_text("CREATE EXTENSION IF NOT EXISTS vector"))
                Base.metadata.create_all(bind=s.get_bind())

                content_sha = hashlib.sha256(data).hexdigest()
                doc = Document(filename=filename, content_type=content_type, source_path=source_path, content_sha256=content_sha)
                s.add(doc)
                s.flush()
                start = 0
                for idx, (ch, emb) in enumerate(zip(chunks, embeddings)):
                    end = start + len(ch)
                    s.add(Chunk(document_id=doc.id, ordinal=idx, text=ch, embedding=emb, start_char=start, end_char=end))
                    start = end - 100 if end - 100 > 0 else end
                s.flush()
                
                span.set_attributes({
                    "rag.document_id": doc.id,
                    "rag.content_sha256": content_sha,
                })
                return doc.id


